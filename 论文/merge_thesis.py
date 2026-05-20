"""
Merge 封面与声明.docx + 正文.docx → 终稿.docx
Handles section-level header/footer/page-number config post-merge.
"""
import copy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"F:\bishe\论文"
FONT_SONG = "SimSun"
FONT_HEI = "SimHei"
LINE_SPACING_PT = 20  # 固定行距 20pt

def set_font(run, name, size_pt, bold=False):
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)

def add_fld_char(para, fld_char_type):
    run = para.add_run()
    el = OxmlElement('w:fldChar')
    el.set(qn('w:fldCharType'), fld_char_type)
    run._element.append(el)

def add_instr_text(para, text):
    run = para.add_run()
    el = OxmlElement('w:instrText')
    el.set(qn('xml:space'), 'preserve')
    el.text = text
    run._element.append(el)

def add_page_field(para, format_type='decimal'):
    add_fld_char(para, 'begin')
    instr = ' PAGE \\* ROMAN ' if format_type == 'roman' else ' PAGE '
    add_instr_text(para, instr)
    add_fld_char(para, 'separate')
    # Default visible text between separate and end (Word replaces with actual page num)
    run = para.add_run()
    run.text = '1'
    add_fld_char(para, 'end')

def add_paragraph_border(para):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def configure_merged_sections(doc):
    """Post-merge: configure headers, footers, page numbers.

    Merged document sections:
      0: Cover        — no header/footer
      1: Declaration  — no header/footer
      2: Copyright    — no header/footer
      3: Front matter — no header, Roman page nums
      4: Main body    — header + Arabic page nums from 1
    """
    sections = doc.sections
    print(f"Total sections: {len(sections)}")

    if len(sections) < 5:
        print("WARNING: expected >=5 sections, section config may be wrong")
        return

    # --- Section 3: Front matter (abstracts + TOC) ---
    sec_fm = sections[3]
    # Header: empty
    hdr = sec_fm.header
    hdr.is_linked_to_previous = False
    # Footer: Roman numeral page numbers, centered, start at I
    ftr = sec_fm.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp, format_type='roman')
    # Roman numerals start at I
    sectPr = sec_fm._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)

    # --- Section 4: Main body ---
    sec_body = sections[4]
    # Header: school name + underline
    hdr = sec_body.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run('应急管理大学本科毕业设计（论文）')
    set_font(run, FONT_SONG, 9)
    add_paragraph_border(hp)

    # Footer: Arabic page numbers, restart at 1
    ftr = sec_body.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp, format_type='decimal')
    # Restart page numbering
    sectPr = sec_body._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)

    # Ensure all sections use A4 page size
    for sec in sections:
        sectPr = sec._sectPr
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is None:
            pgSz = OxmlElement('w:pgSz')
            sectPr.insert(0, pgSz)
        pgSz.set(qn('w:w'), '11906')
        pgSz.set(qn('w:h'), '16838')

    print("Section headers/footers configured.")


def _fix_merged_styles(doc):
    """Fix Normal + Heading 1-3 styles in merged document to match template.

    Called post-merge because the merge process inherits the cover document's
    default python-docx styles, which don't match the school template.
    """
    def _ensure(el, tag):
        child = el.find(qn('w:' + tag))
        if child is None:
            child = OxmlElement('w:' + tag)
            el.insert(0, child)
        return child

    def _set_rfonts(rPr, ascii_f, ea_f, hAnsi_f=None):
        if hAnsi_f is None:
            hAnsi_f = ascii_f
        rFonts = _ensure(rPr, 'rFonts')
        rFonts.set(qn('w:ascii'), ascii_f)
        rFonts.set(qn('w:eastAsia'), ea_f)
        rFonts.set(qn('w:hAnsi'), hAnsi_f)

    def _set_color(rPr, hex_color):
        color = rPr.find(qn('w:color'))
        if color is None:
            color = OxmlElement('w:color')
            rPr.append(color)
        color.set(qn('w:val'), hex_color)

    # --- Normal (may not exist in docx-js generated cover) ---
    try:
        normal = doc.styles['Normal']
    except KeyError:
        # Create Normal style by adding to styles element
        styles_el = doc.styles.element
        nsmap = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        ns = '{' + nsmap + '}'
        style_el = OxmlElement('w:style')
        style_el.set(ns + 'type', 'paragraph')
        style_el.set(ns + 'styleId', 'Normal')
        style_el.set(ns + 'default', '1')
        name_el = OxmlElement('w:name')
        name_el.set(ns + 'val', 'Normal')
        style_el.append(name_el)
        styles_el.append(style_el)
        normal = doc.styles['Normal']

    normal.font.name = FONT_SONG
    normal.font.size = Pt(12)
    normal.font.color.rgb = None  # reset
    normal.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # eastAsia via XML
    nrPr = _ensure(normal.element, 'rPr')
    _set_rfonts(nrPr, FONT_SONG, FONT_SONG)
    _set_color(nrPr, '000000')

    # --- Headings ---
    specs = {
        'Heading1': {'sz': '30', 'before': '400', 'after': '400', 'jc': 'center'},
        'Heading2': {'sz': '28', 'before': '480', 'after': '120', 'jc': 'left'},
        'Heading3': {'sz': '24', 'before': '240', 'after': '120', 'jc': 'left'},
    }
    for name, spec in specs.items():
        hs = doc.styles[name]
        hs.font.name = FONT_HEI
        hs.font.size = Pt(int(spec['sz']) / 2)
        hs.font.bold = True
        hs.font.color.rgb = None  # reset (removes blue)
        hs.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
        hs.paragraph_format.space_before = Pt(int(spec['before']) / 20)
        hs.paragraph_format.space_after = Pt(int(spec['after']) / 20)
        hs.paragraph_format.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if spec['jc'] == 'center'
            else WD_ALIGN_PARAGRAPH.LEFT
        )
        # XML-level fixes
        hrPr = _ensure(hs.element, 'rPr')
        _set_rfonts(hrPr, FONT_HEI, FONT_HEI)
        _set_color(hrPr, '000000')
        # paragraph properties
        hpPr = _ensure(hs.element, 'pPr')
        sp = _ensure(hpPr, 'spacing')
        sp.set(qn('w:before'), spec['before'])
        sp.set(qn('w:after'), spec['after'])
        sp.set(qn('w:line'), str(LINE_SPACING_PT * 20))
        sp.set(qn('w:lineRule'), 'exact')
        jc = _ensure(hpPr, 'jc')
        jc.set(qn('w:val'), spec['jc'])
        # No first-line indent
        ind = _ensure(hpPr, 'ind')
        ind.set(qn('w:firstLine'), '0')

    print("Merged document styles fixed.")


# === Main ===
print("Loading cover...")
cover = Document(r'F:\bishe\论文\封面与声明.docx')
print(f"Cover sections: {len(cover.sections)}")

print("Loading body...")
body = Document(r'F:\bishe\论文\正文.docx')
print(f"Body sections: {len(body.sections)}")
print(f"Body paragraphs: {len(body.paragraphs)}")

# Merge: insert page break, then deep-copy all body elements
print("Merging...")
# Page break to separate declaration from abstract
pb_para = OxmlElement('w:p')
pb_run = OxmlElement('w:r')
pb_br = OxmlElement('w:br')
pb_br.set(qn('w:type'), 'page')
pb_run.append(pb_br)
pb_para.append(pb_run)
cover.element.body.append(copy.deepcopy(pb_para))

for bp in body.element.body:
    el = copy.deepcopy(bp)
    cover.element.body.append(el)

print(f"Post-merge body children: {len(cover.element.body)}")

# Save intermediate, re-open to get correct section count
tmp = r'F:\bishe\论文\_tmp_merged.docx'
cover.save(tmp)

print("Re-opening merged to configure sections...")
merged = Document(tmp)
print(f"Merged sections: {len(merged.sections)}")

configure_merged_sections(merged)
_fix_merged_styles(merged)

output = r'F:\bishe\论文\终稿-标准格式.docx'
merged.save(output)
print(f"\nDone: {output}")

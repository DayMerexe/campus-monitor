"""
组装毕设论文终稿.docx
封面 由 generate_cover.cjs 生成（封面与声明.docx）
正文 从 .md 文件转换

格式依据：学校模板_原版.docx styles.xml
  - 正文：宋体/Times New Roman 12pt，固定行距20pt，首行缩进2字符，两端对齐
  - H1：黑体 15pt，居中，段前40pt 段后20pt
  - H2：黑体 14pt，左对齐，段前24pt 段后6pt
  - H3：黑体 13pt，左对齐，段前12pt 段后6pt
  - 页眉：小五宋体，底部边框线
  - 页脚：小五，居中页码
"""
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = r"F:\bishe\论文"

# --- Font constants ---
FONT_HEI = "SimHei"    # 黑体 — titles
FONT_SONG = "SimSun"   # 宋体 — body
FONT_FANG = "FangSong" # 仿宋 — cover fields
FONT_KAITI = "KaiTi"   # 楷体 — abstract labels

# --- Template spec constants (in twips, 1pt = 20 twips) ---
LINE_SPACING_PT = 20       # 固定行距 20pt (= 400 twips)
FIRST_LINE_INDENT_TWIPS = 480  # 2 chars × 12pt × 20


def _ensure_rPr(element):
    """Ensure a w:rPr child exists, return it."""
    rPr = element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        element.insert(0, rPr)
    return rPr


def _set_rFonts(rPr, ascii_font, eastasia_font, hAnsi_font=None):
    """Set w:rFonts attributes on a run-properties element."""
    if hAnsi_font is None:
        hAnsi_font = ascii_font
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:eastAsia'), eastasia_font)
    rFonts.set(qn('w:hAnsi'), hAnsi_font)


def _set_style_eastasia(style, font_name):
    """Set east-Asian font on a style's rPr (python-docx font.name doesn't)."""
    rPr = _ensure_rPr(style.element)
    _set_rFonts(rPr, font_name, font_name)


def _set_style_color(style, hex_color):
    """Set font color on a style's rPr."""
    rPr = _ensure_rPr(style.element)
    color = rPr.find(qn('w:color'))
    if color is None:
        color = OxmlElement('w:color')
        rPr.append(color)
    color.set(qn('w:val'), hex_color)


def set_font(run, name, size_pt, bold=False, color='000000'):
    """Set run font with ascii + east-asian + color."""
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = name
    rPr = _ensure_rPr(run._element)
    _set_rFonts(rPr, name, name)
    # Ensure color (override any style-inherited color like blue headings)
    c = rPr.find(qn('w:color'))
    if c is None:
        c = OxmlElement('w:color')
        rPr.append(c)
    c.set(qn('w:val'), color)


def set_font_eng(run, size_pt, bold=False):
    """Set English font (Times New Roman) with correct eastAsia fallback."""
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = 'Times New Roman'
    rPr = _ensure_rPr(run._element)
    _set_rFonts(rPr, 'Times New Roman', '宋体')


def set_line_spacing_exact(para, line_pt=LINE_SPACING_PT):
    """Set exact (fixed) line spacing in points. Template uses 20pt exact."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(line_pt * 20)))
    spacing.set(qn('w:lineRule'), 'exact')


def set_first_line_indent(para, indent_chars=2, font_size_pt=12):
    """Set first line indent.

    w:firstLine uses twips (1/20 of a point).
    2 chars at 12pt = 24pt = 480 twips.
    """
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLine'), str(font_size_pt * indent_chars * 20))


def add_body_para(doc, text, first_indent=True, font_size=12, line_spacing=LINE_SPACING_PT, space_before=0):
    """Add a body paragraph with template-compliant formatting.

    Default: 小四(12pt) 宋体, 行距20磅. For references: 五号(10.5pt), 行距16磅, 段前3磅.
    """
    p = doc.add_paragraph()
    if first_indent:
        set_first_line_indent(p, font_size_pt=font_size)
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    set_line_spacing_exact(p, line_pt=line_spacing)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, FONT_SONG, font_size)
    return p


def add_heading_styled(doc, text, level):
    """Add a chapter/section heading matching template specs.

    Template (per TEXTBOX 17-19):
      H1: 黑体小三(15pt), centered, before=20pt after=20pt
      H2: 黑体四号(14pt), left, before=24pt after=6pt
      H3: 黑体小四(12pt), left, before=12pt after=6pt
    """
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    set_line_spacing_exact(p)

    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(20)
        size = 15
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(6)
        size = 14
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        size = 12

    run = p.add_run(text)
    set_font(run, FONT_HEI, size, bold=True, color='000000')
    return p


def add_toc(doc):
    """Insert TOC page, then a section break separating front matter from body."""
    # "目录" title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing_exact(p)
    run = p.add_run('目录')
    set_font(run, FONT_HEI, 15, bold=True, color='000000')

    # TOC field code
    toc_p = doc.add_paragraph()
    _add_fld_char(toc_p, 'begin')
    _add_instr_text(toc_p, ' TOC \\o "1-3" \\h \\z \\u ')
    _add_fld_char(toc_p, 'separate')
    run = toc_p.add_run('（请在Word中右键此区域 → 更新域 → 更新整个目录）')
    set_font(run, FONT_SONG, 10)
    _add_fld_char(toc_p, 'end')

    # Section break: front matter ends, main body begins
    doc.add_section()


def _add_fld_char(para, fld_char_type):
    """Helper: add a w:fldChar element to a paragraph"""
    run = para.add_run()
    el = OxmlElement('w:fldChar')
    el.set(qn('w:fldCharType'), fld_char_type)
    run._element.append(el)


def _add_instr_text(para, text):
    """Helper: add a w:instrText element to a paragraph"""
    run = para.add_run()
    el = OxmlElement('w:instrText')
    el.set(qn('xml:space'), 'preserve')
    el.text = text
    run._element.append(el)


def add_table_from_md(doc, lines, start_idx):
    """Parse and add a markdown table. Returns (table_element, end_line_idx)"""
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1

    if len(table_lines) < 2:
        return None, start_idx

    rows = []
    for tl in table_lines:
        cells = [c.strip() for c in tl.strip('|').split('|')]
        if all(re.match(r'^[-:]+$', c.strip()) for c in cells):
            continue
        rows.append(cells)

    if not rows:
        return None, i - 1

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols, style='Table Grid')

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                is_header = (row_idx == 0)
                set_font(run, FONT_HEI if is_header else FONT_SONG, 9)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if is_header:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'D9E2F3')
                    shading.set(qn('w:val'), 'clear')
                    cell._element.find(qn('w:tcPr')).append(shading)

    return table, i - 1


def add_abstract_section(doc, md_path, is_english=False):
    """Add abstract with template-compliant formatting"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    if is_english:
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('# '):
                continue
            if line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line.replace('**', ''))
                set_font_eng(run, 12, bold=True)
                set_line_spacing_exact(p)
            else:
                p = doc.add_paragraph()
                set_first_line_indent(p)
                run = p.add_run(line)
                set_font_eng(run, 12)
                set_line_spacing_exact(p)
    else:
        heading_found = False
        for line in lines:
            line = line.strip()
            if not line and not heading_found:
                continue
            if line.startswith('# '):
                add_heading_styled(doc, line[2:], 1)
                heading_found = True
                continue
            if line.startswith('**') and line.endswith('**') and '关键词' in line:
                p = doc.add_paragraph()
                run = p.add_run(line.replace('**', ''))
                set_font(run, FONT_HEI, 12, bold=True)
                set_line_spacing_exact(p)
            elif line:
                cleaned = line.lstrip('　 ')
                add_body_para(doc, cleaned)

    doc.add_page_break()


def parse_md_body(doc, md_path, body_font_size=12, body_line_spacing=LINE_SPACING_PT, body_space_before=0, body_indent=True):
    """Parse a markdown chapter file and add to document.

    body_font_size/body_line_spacing overrides for special sections (e.g. 参考文献 10.5pt/16pt).
    """
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Skip mermaid code blocks
        if stripped.startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1
            continue

        # Skip HTML tags (figure placeholders)
        if stripped.startswith('<p align') or stripped.startswith('<b>图') or stripped.startswith('</p>'):
            i += 1
            continue

        # Headings
        if stripped.startswith('# '):
            text = stripped[2:]
            add_heading_styled(doc, text, 1)
            i += 1
            continue
        elif stripped.startswith('## '):
            text = stripped[3:]
            add_heading_styled(doc, text, 2)
            i += 1
            continue
        elif stripped.startswith('### '):
            text = stripped[4:]
            add_heading_styled(doc, text, 3)
            i += 1
            continue

        # Tables
        if stripped.startswith('|'):
            table, next_i = add_table_from_md(doc, lines, i)
            if table:
                p = doc.add_paragraph()
                i = next_i + 1
                continue

        cleaned = stripped.lstrip('　 ')

        # Bold leading text with **...**
        if cleaned.startswith('**') and '**' in cleaned[2:]:
            p = doc.add_paragraph()
            set_first_line_indent(p)
            set_line_spacing_exact(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            segments = re.split(r'(\*\*[^*]+\*\*)', cleaned)
            for seg in segments:
                if seg.startswith('**') and seg.endswith('**'):
                    run = p.add_run(seg[2:-2])
                    set_font(run, FONT_HEI, 12, bold=True)
                elif seg:
                    run = p.add_run(seg)
                    set_font(run, FONT_SONG, 12)
            i += 1
            continue

        # Table captions (like "表2-1 xxx") — 黑体11磅
        if re.match(r'^表\d', cleaned):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cleaned)
            set_font(run, FONT_HEI, 11, bold=True)
            i += 1
            continue

        # Figure captions (like "图2-1 xxx") — 宋体11磅
        if re.match(r'^图\d', cleaned):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cleaned)
            set_font(run, FONT_SONG, 11)
            i += 1
            continue

        # Formula or equation
        if cleaned.startswith('CCR1') or cleaned.startswith('饱和度和i') or cleaned.startswith('式（'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_line_spacing_exact(p)
            run = p.add_run(cleaned)
            set_font(run, FONT_SONG, 12)
            i += 1
            continue

        # Normal body paragraph
        add_body_para(doc, cleaned, first_indent=body_indent, font_size=body_font_size,
                      line_spacing=body_line_spacing, space_before=body_space_before)
        i += 1

    doc.add_page_break()


def configure_sections(doc):
    """Setup headers, footers, and page numbers.

    Section 0: Front matter — no header, Roman numeral page numbers
    Section 1: Main body — header with school name + underline, Arabic page numbers from 1
    """
    _setup_front_matter(doc.sections[0])
    _setup_body_section(doc.sections[1])


def _setup_front_matter(section):
    """No header, Roman numeral page numbers centered in footer."""
    section.different_first_page_header_footer = False
    header = section.header
    header.is_linked_to_previous = False
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_field(fp, format_type='roman')


def _setup_body_section(section):
    """Header: school name + bottom border. Footer: Arabic page numbers from 1."""
    section.different_first_page_header_footer = False

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run('应急管理大学本科毕业设计（论文）')
    set_font(run, FONT_SONG, 9)
    _add_paragraph_border(hp)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_field(fp, format_type='decimal')
    sectPr = section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)


def _add_page_field(para, format_type='decimal'):
    """Insert a PAGE field code."""
    _add_fld_char(para, 'begin')
    instr = ' PAGE \\* ROMAN ' if format_type == 'roman' else ' PAGE '
    _add_instr_text(para, instr)
    _add_fld_char(para, 'separate')
    _add_fld_char(para, 'end')


def _add_paragraph_border(para):
    """Add thin bottom border to a paragraph (header underline)."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_page_to_a4(section):
    """Set section page size to A4 (11906 × 16838 DXA)."""
    sectPr = section._sectPr
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz')
        sectPr.insert(0, pgSz)
    pgSz.set(qn('w:w'), '11906')
    pgSz.set(qn('w:h'), '16838')


def main():
    global doc
    doc = Document()

    # ============================================================
    # 1. Fix Normal style to match template
    #    Template: 宋体/Times New Roman 12pt, 固定行距20pt, 两端对齐
    # ============================================================
    style = doc.styles['Normal']
    style.font.name = FONT_SONG
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)  # exact 20pt
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_style_eastasia(style, FONT_SONG)
    _set_style_color(style, '000000')

    # ============================================================
    # 2. Fix Heading styles (1-3) via direct XML manipulation
    #    python-docx's high-level API doesn't set eastAsia or color
    # ============================================================
    heading_specs = {
        1: {'sz': '30', 'before': '400', 'after': '400', 'jc': 'center'},   # 小三 15pt
        2: {'sz': '28', 'before': '480', 'after': '120', 'jc': 'left'},      # 四号 14pt
        3: {'sz': '24', 'before': '240', 'after': '120', 'jc': 'left'},      # 小四 12pt
    }
    for level, spec in heading_specs.items():
        hs = doc.styles[f'Heading {level}']

        # font
        hs.font.name = FONT_HEI
        hs.font.size = Pt(int(spec['sz']) / 2)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)

        # paragraph format
        hs.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
        hs.paragraph_format.space_before = Pt(int(spec['before']) / 20)
        hs.paragraph_format.space_after = Pt(int(spec['after']) / 20)
        hs.paragraph_format.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if spec['jc'] == 'center'
            else WD_ALIGN_PARAGRAPH.LEFT
        )

        # east-Asian font (python-docx doesn't set this)
        _set_style_eastasia(hs, FONT_HEI)
        _set_style_color(hs, '000000')

        # ensure style paragraph properties include spacing and alignment
        _fix_style_pPr(hs, spec)

    # ============================================================
    # 3. Force A4 page size on first section
    # ============================================================
    _set_page_to_a4(doc.sections[0])

    # ============================================================
    # Assemble content
    # ============================================================
    # 中文摘要
    add_abstract_section(doc, os.path.join(BASE, '中文摘要.md'), is_english=False)
    # 英文摘要
    add_abstract_section(doc, os.path.join(BASE, '英文摘要.md'), is_english=True)
    # 目录 (adds section break internally)
    add_toc(doc)

    # Set A4 on the new section (body) too
    if len(doc.sections) > 1:
        _set_page_to_a4(doc.sections[1])

    # 正文五章
    chapters = [
        '第1章-绪论.md',
        '第2章-需求分析.md',
        '第3章-系统设计.md',
        '第4章-详细设计与实现.md',
        '第5章-测试.md',
    ]
    for ch_file in chapters:
        path = os.path.join(BASE, ch_file)
        if os.path.exists(path):
            print(f"Processing: {ch_file}")
            parse_md_body(doc, path)
        else:
            print(f"WARNING: {ch_file} not found")

    # 结论
    conclusion_path = os.path.join(BASE, '结论.md')
    if os.path.exists(conclusion_path):
        print("Processing: 结论.md")
        parse_md_body(doc, conclusion_path)

    # 参考文献 — 五号(10.5pt), 行距16磅, 段前3磅 (TEXTBOX 36)
    ref_path = os.path.join(BASE, '参考文献.md')
    if os.path.exists(ref_path):
        print("Processing: 参考文献.md")
        parse_md_body(doc, ref_path, body_font_size=10.5, body_line_spacing=16,
                      body_space_before=3, body_indent=False)

    # 致谢
    thanks_path = os.path.join(BASE, '致谢.md')
    if os.path.exists(thanks_path):
        print("Processing: 致谢.md")
        parse_md_body(doc, thanks_path)

    # 附录
    appendix_path = os.path.join(BASE, '附录.md')
    if os.path.exists(appendix_path):
        print("Processing: 附录.md")
        parse_md_body(doc, appendix_path)

    output = os.path.join(BASE, '正文.docx')
    doc.save(output)
    print(f"\nDone: {output}")


def _fix_style_pPr(style, spec):
    """Ensure the style's pPr has spacing and alignment set at XML level.

    python-docx's paragraph_format sometimes doesn't serialize correctly
    for styles, so we set them directly on the XML element.
    """
    pPr = style.element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        style.element.append(pPr)

    # Spacing
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), spec['before'])
    spacing.set(qn('w:after'), spec['after'])
    spacing.set(qn('w:line'), str(LINE_SPACING_PT * 20))
    spacing.set(qn('w:lineRule'), 'exact')

    # Justification
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), spec['jc'])

    # No first-line indent for headings
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLine'), '0')


if __name__ == '__main__':
    main()
